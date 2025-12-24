# 代碼功能說明: DOCX 文件解析器
# 創建日期: 2025-01-27 23:30 (UTC+8)
# 創建人: Daniel Chung
# 最後修改日期: 2025-01-27 23:30 (UTC+8)

"""DOCX 文件解析器 - 使用 python-docx"""

from typing import Any, Dict

from .base_parser import BaseParser

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocxParser(BaseParser):
    """DOCX 文件解析器"""

    def __init__(self):
        super().__init__()
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 未安裝，請運行: pip install python-docx")

    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        解析 DOCX 文件

        Args:
            file_path: 文件路徑

        Returns:
            解析結果，包含文本內容和結構元數據
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 未安裝")

        try:
            doc = Document(file_path)

            # 提取段落文本
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # 提取表格
            tables_data = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append(
                    {
                        "table_index": table_idx,
                        "rows": len(table_data),
                        "data": table_data,
                    }
                )

            full_text = "\n".join(paragraphs)

            return {
                "text": full_text,
                "metadata": {
                    "num_paragraphs": len(paragraphs),
                    "num_tables": len(tables_data),
                    "tables": tables_data,
                    "char_count": len(full_text),
                },
            }
        except Exception as e:
            self.logger.error("DOCX 文件解析失敗", file_path=file_path, error=str(e))
            raise

    def parse_from_bytes(self, file_content: bytes, **kwargs) -> Dict[str, Any]:
        """
        從字節內容解析 DOCX

        Args:
            file_content: 文件內容（字節）
            **kwargs: 其他參數

        Returns:
            解析結果
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 未安裝")

        try:
            from io import BytesIO

            doc = Document(BytesIO(file_content))

            # 提取段落文本
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # 提取表格
            tables_data = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append(
                    {
                        "table_index": table_idx,
                        "rows": len(table_data),
                        "data": table_data,
                    }
                )

            full_text = "\n".join(paragraphs)

            return {
                "text": full_text,
                "metadata": {
                    "num_paragraphs": len(paragraphs),
                    "num_tables": len(tables_data),
                    "tables": tables_data,
                    "char_count": len(full_text),
                },
            }
        except Exception as e:
            self.logger.error("DOCX 解析失敗", error=str(e))
            raise

    def get_supported_extensions(self) -> list:
        return [".docx"]
