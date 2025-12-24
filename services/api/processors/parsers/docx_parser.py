# 代碼功能說明: DOCX 文件解析器
# 創建日期: 2025-12-06
# 創建人: Daniel Chung
# 最後修改日期: 2025-12-20

"""DOCX 文件解析器 - 使用 python-docx 和 python-mammoth"""

import os
from typing import Any, Dict, List, Optional

from .base_parser import BaseParser

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import mammoth  # noqa: F401

    MAMMOTH_AVAILABLE = True
except ImportError:
    MAMMOTH_AVAILABLE = False


class DocxParser(BaseParser):
    """DOCX 文件解析器"""

    def __init__(self):
        super().__init__()
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 未安裝，請運行: pip install python-docx")

    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        解析 DOCX 文件

        Args:
            file_path: 文件路徑

        Returns:
            解析結果，包含文本內容和結構元數據
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 未安裝")

        try:
            doc = Document(file_path)

            # 提取段落文本
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # 提取表格
            tables_data = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append(
                    {
                        "table_index": table_idx,
                        "rows": len(table_data),
                        "data": table_data,
                    }
                )

            full_text = "\n".join(paragraphs)

            return {
                "text": full_text,
                "metadata": {
                    "num_paragraphs": len(paragraphs),
                    "num_tables": len(tables_data),
                    "tables": tables_data,
                    "char_count": len(full_text),
                },
            }
        except Exception as e:
            self.logger.error("DOCX 文件解析失敗", file_path=file_path, error=str(e))
            raise

    def parse_from_bytes(self, file_content: bytes, **kwargs) -> Dict[str, Any]:
        """
        從字節內容解析 DOCX

        Args:
            file_content: 文件內容（字節）
            **kwargs: 其他參數

        Returns:
            解析結果
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 未安裝")

        try:
            from io import BytesIO

            doc = Document(BytesIO(file_content))

            # 提取段落文本
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # 提取表格
            tables_data = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append(
                    {
                        "table_index": table_idx,
                        "rows": len(table_data),
                        "data": table_data,
                    }
                )

            full_text = "\n".join(paragraphs)

            return {
                "text": full_text,
                "metadata": {
                    "num_paragraphs": len(paragraphs),
                    "num_tables": len(tables_data),
                    "tables": tables_data,
                    "char_count": len(full_text),
                },
            }
        except Exception as e:
            self.logger.error("DOCX 解析失敗", error=str(e))
            raise

    def parse_to_markdown(
        self,
        file_path: str,
        output_dir: Optional[str] = None,
        extract_images: bool = True,
    ) -> Dict[str, Any]:
        """
        將 DOCX 轉換為 Markdown 格式

        Args:
            file_path: DOCX 文件路徑
            output_dir: 圖片輸出目錄（可選）
            extract_images: 是否提取圖片

        Returns:
            轉換結果，包含 markdown 文本、元數據和映射表
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 未安裝")

        try:
            from ..offset_mapper import OffsetMapper

            original_text = ""
            images: List[Dict[str, Any]] = []
            tables: List[Dict[str, Any]] = []

            doc = Document(file_path)

            # 提取原始文本
            paragraphs_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs_text.append(para.text)
            original_text = "\n".join(paragraphs_text)

            # 使用 python-docx 提取內容並轉換為 Markdown
            markdown_lines: List[str] = []

            # 處理段落
            for para in doc.paragraphs:
                if not para.text.strip():
                    markdown_lines.append("")
                    continue

                # 檢查段落樣式
                style_name = para.style.name if para.style else ""
                text = para.text

                # 映射樣式到 Markdown
                markdown_line = self._map_style_to_markdown(style_name, text)
                markdown_lines.append(markdown_line)

            # 處理表格
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)

                if table_data:
                    markdown_table = self._table_to_markdown(table_data)
                    tables.append(
                        {
                            "table_index": table_idx,
                            "markdown": markdown_table,
                        }
                    )
                    markdown_lines.append("")
                    markdown_lines.append(markdown_table)
                    markdown_lines.append("")

            # 提取圖片（如果啟用）
            if extract_images and output_dir:
                images = self._extract_images(doc, file_path, output_dir)

            markdown_text = "\n".join(markdown_lines)

            # 創建偏移量映射
            mapper = OffsetMapper()
            offset_mapping = mapper.create_mapping(original_text, markdown_text)

            return {
                "text": markdown_text,
                "metadata": {
                    "original_text": original_text,
                    "num_paragraphs": len(doc.paragraphs),
                    "num_tables": len(tables),
                    "tables": tables,
                    "images": images,
                    "offset_mapping": offset_mapping,
                    "mapping_stats": mapper.get_mapping_stats(),
                },
            }
        except Exception as e:
            self.logger.error("DOCX 轉 Markdown 失敗", file_path=file_path, error=str(e))
            raise

    def _map_style_to_markdown(self, style_name: str, text: str) -> str:
        """
        將 Word 樣式映射到 Markdown 格式

        Args:
            style_name: Word 樣式名稱
            text: 文本內容

        Returns:
            Markdown 格式的文本
        """
        style_lower = style_name.lower()

        # 標題樣式映射
        if "heading 1" in style_lower or "title" in style_lower:
            return f"# {text}"
        elif "heading 2" in style_lower:
            return f"## {text}"
        elif "heading 3" in style_lower:
            return f"### {text}"
        elif "heading 4" in style_lower:
            return f"#### {text}"
        elif "heading 5" in style_lower:
            return f"##### {text}"
        elif "heading 6" in style_lower:
            return f"###### {text}"
        elif "list" in style_lower or "bullet" in style_lower:
            return f"- {text}"
        elif "number" in style_lower:
            # 簡單處理，實際可能需要更複雜的邏輯
            return f"1. {text}"
        else:
            # 普通段落
            return text

    def _table_to_markdown(self, table_data: List[List[str]]) -> str:
        """
        將表格數據轉換為 Markdown 表格格式

        Args:
            table_data: 表格數據（二維列表）

        Returns:
            Markdown 表格字符串
        """
        if not table_data:
            return ""

        markdown_lines: List[str] = []

        # 確保所有行的長度一致
        max_cols = max(len(row) for row in table_data) if table_data else 0
        normalized_table = []
        for row in table_data:
            normalized_row = row + [""] * (max_cols - len(row))
            normalized_row = normalized_row[:max_cols]
            normalized_table.append(normalized_row)

        if not normalized_table:
            return ""

        # 生成表頭
        header = normalized_table[0]
        markdown_lines.append("| " + " | ".join(header) + " |")
        markdown_lines.append("| " + " | ".join(["---"] * len(header)) + " |")

        # 生成數據行
        for row in normalized_table[1:]:
            markdown_lines.append("| " + " | ".join(row) + " |")

        return "\n".join(markdown_lines)

    def _extract_images(self, doc: Any, file_path: str, output_dir: str) -> List[Dict[str, Any]]:
        """
        從 DOCX 文件中提取圖片

        Args:
            doc: Document 對象
            file_path: 原始文件路徑
            output_dir: 圖片輸出目錄

        Returns:
            圖片信息列表
        """
        images: List[Dict[str, Any]] = []

        if not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)

        try:
            # python-docx 不直接支持圖片提取
            # 需要從 DOCX 的內部結構中提取
            # 這裡提供一個簡化的實現
            # 實際應用中可能需要使用 zipfile 解壓 DOCX 並提取圖片

            # 簡化實現：返回空列表
            # 完整實現需要：
            # 1. 解壓 DOCX（實際是 ZIP 文件）
            # 2. 從 word/media/ 目錄提取圖片
            # 3. 保存到 output_dir
            # 4. 生成 Markdown 圖片引用

            self.logger.info("圖片提取功能待實現", file_path=file_path)
        except Exception as e:
            self.logger.warning("圖片提取失敗", error=str(e))

        return images

    def get_supported_extensions(self) -> list:
        return [".docx"]
